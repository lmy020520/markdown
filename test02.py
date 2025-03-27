import os
from docx import Document
import PyPDF2
import openpyxl

def docx_to_md(input_path, output_path):
    doc = Document(input_path)
    with open(output_path, 'w', encoding='utf-8') as md_file:
        for para in doc.paragraphs:
            text = para.text
            # 清理字符串，移除无法编码的字符
            clean_text = text.encode('utf-8', errors='ignore').decode('utf-8')
            # 检测标题样式并转换为 Markdown 标题
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading', ''))
                md_file.write(f"{'#' * level} {clean_text}\n\n")
            else:
                md_file.write(f"{clean_text}\n\n")
        # 处理表格
        for table in doc.tables:
            # 写入表格分隔线
            headers = ['---' for _ in range(len(table.columns))]
            md_file.write('| ' + ' | '.join(headers) + ' |\n')
            # 写入表头和内容
            for row in table.rows:
                cells = [cell.text.strip().encode('utf-8', errors='ignore').decode('utf-8') for cell in row.cells]
                md_file.write('| ' + ' | '.join(cells) + ' |\n')
            md_file.write('\n')

def pdf_to_md(input_path, output_path):
    with open(input_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        with open(output_path, 'w', encoding='utf-8') as md_file:
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    # 清理字符串，移除无法编码的字符
                    clean_text = text.encode('utf-8', errors='ignore').decode('utf-8')
                    md_file.write(clean_text + "\n\n")
                else:
                    md_file.write(f"Page {page_num + 1} has no extractable text.\n\n")

def txt_to_md(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as txt_file:
        with open(output_path, 'w', encoding='utf-8') as md_file:
            for line in txt_file:
                # 清理字符串，移除无法编码的字符
                clean_line = line.encode('utf-8', errors='ignore').decode('utf-8')
                md_file.write(clean_line)

def xlsx_to_md(input_path, output_path):
    workbook = openpyxl.load_workbook(input_path)
    sheet = workbook.active
    with open(output_path, 'w', encoding='utf-8') as md_file:
        # 写入表头
        headers = [str(cell.value).encode('utf-8', errors='ignore').decode('utf-8') for cell in sheet[1]]
        md_file.write('| ' + ' | '.join(headers) + ' |\n')
        # 写入表格分隔线
        md_file.write('| ' + ' | '.join(['---' for _ in headers]) + ' |\n')
        # 写入数据行
        for row in sheet.iter_rows(min_row=2, values_only=True):
            cleaned_row = [str(value).encode('utf-8', errors='ignore').decode('utf-8') for value in row]
            md_file.write('| ' + ' | '.join(cleaned_row) + ' |\n')
        md_file.write('\n')

def convert_file(input_path, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    file_name = os.path.basename(input_path)
    base_name, ext = os.path.splitext(file_name)
    output_path = os.path.join(output_dir, f"{base_name}.md")
    
    if ext.lower() == '.docx':
        docx_to_md(input_path, output_path)
    elif ext.lower() == '.pdf':
        pdf_to_md(input_path, output_path)
    elif ext.lower() == '.txt':
        txt_to_md(input_path, output_path)
    elif ext.lower() == '.xlsx':
        xlsx_to_md(input_path, output_path)
    else:
        print(f"Unsupported file format: {ext}")
        return
    
    print(f"Converted {file_name} to {output_path}")

if __name__ == "__main__":
    # 示例用法
    input_files = [
        # "example.docx",
        "pdf/lmy.pdf",
        # "example.txt",
        # "example.xlsx"
    ]
    
    output_directory = "res"
    
    for file_path in input_files:
        if os.path.exists(file_path):
            convert_file(file_path, output_directory)
        else:
            print(f"File not found: {file_path}")