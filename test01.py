import os
import pandas as pd
import pdfplumber
from docx import Document


class MarkdownifyMCP:
    def convert(self, input_path, output_dir):
        """主转换方法"""
        ext = os.path.splitext(input_path)[1].lower()
        
        if ext == '.docx':
            return self._convert_docx(input_path, output_dir)
        elif ext == '.pdf':
            return self._convert_pdf(input_path, output_dir)
        elif ext in ('.xlsx', '.xls'):
            return self._convert_excel(input_path, output_dir)
        else:
            raise ValueError("Unsupported file format")

    def _convert_docx(self, docx_path, output_dir):
        """处理 Word 文档"""
        doc = Document(docx_path)
        md_content = []
        
        for para in doc.paragraphs:
            # 标题检测
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split(' ')[1])
                md_content.append(f"{'#' * level} {para.text}")
            else:
                md_content.append(para.text)
        
        # 表格处理
        for table in doc.tables:
            md_table = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                md_table.append("| " + " | ".join(cells) + " |")
            
            # 添加表头分隔符
            if md_table:
                md_table.insert(1, "| " + " | ".join(["---"] * len(table.columns)) + " |")
            
            md_content.extend(["\n", "### 表格提取\n"] + md_table)
        
        self._save_md(md_content, docx_path, output_dir)

    def _convert_pdf(self, pdf_path, output_dir):
        """处理 PDF 文档（带异常捕获）"""
        md_content = []
        
        with pdfplumber.open(pdf_path) as pdf:
            try:
                font_size_map = self._analyze_font_sizes(pdf)
            except Exception as e:
                print(f"字体分析失败: {str(e)}")
                font_size_map = {}

            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    # 获取页面实际尺寸
                    page_width = page.width
                    page_height = page.height

                    # 安全获取文本块
                    text_blocks = page.extract_text_lines(
                        extra_attrs=["fontname", "size"]
                    ) or []

                    for block in text_blocks:
                        try:
                            text = block.get('text', '').strip()
                            if not text:
                                continue

                            # 安全获取布局信息
                            bbox = block.get('bbox', (0, 0, page_width, page_height))
                            chars = block.get('chars', [])
                            
                            # 标题检测增强
                            is_heading = False
                            heading_level = 0

                            if chars:
                                # 安全获取字体尺寸
                                first_char = chars[0]
                                font_size = first_char.get('size', 12)
                                level = font_size_map.get(font_size, 0)
                                
                                # 计算居中状态
                                x_center = (bbox[2] + bbox[0]) / 2
                                is_centered = abs(x_center - page_width/2) < 50
                                
                                # 判断大写
                                is_uppercase = text.isupper()
                                
                                # 综合判断
                                conditions = [level > 0, is_centered, is_uppercase]
                                if sum(conditions) >= 2:
                                    is_heading = True
                                    heading_level = min(level, 6)

                            if is_heading:
                                md_content.append(f"\n{'#' * heading_level} {text}\n")
                            else:
                                md_content.append(text + "  ")

                        except Exception as block_error:
                            print(f"文本块处理失败: {str(block_error)}")
                            continue

                    # 表格处理（带异常捕获）
                    try:
                        tables = page.extract_tables()
                        for table in tables:
                            if table and len(table) > 0:
                                df = pd.DataFrame(table[1:], columns=table[0])
                                md_content.append("\n### PDF表格\n" + df.to_markdown(index=False))
                    except Exception as table_error:
                        print(f"表格处理失败: {str(table_error)}")

                except Exception as page_error:
                    print(f"第 {page_num} 页处理失败: {str(page_error)}")
                    continue

        self._save_md(md_content, pdf_path, output_dir)

    def _analyze_font_sizes(self, pdf):
        """分析全文档字体尺寸分布"""
        font_sizes = []
        for page in pdf.pages:
            chars = page.chars
            if chars:
                font_sizes.extend([char['size'] for char in chars])
        
        # 按字体尺寸排序并划分层级
        unique_sizes = sorted(list(set(font_sizes)), reverse=True)
        return {
            size: min(i+1, 6)  # 最大支持 h6
            for i, size in enumerate(unique_sizes[:5])  # 取前5大尺寸作为标题
        }

    def _detect_heading(self, block, font_size_map):
        """判断是否为标题"""
        # 规则1：字体尺寸匹配
        font_size = block['chars'][0]['size']
        level = font_size_map.get(font_size, 0)
        
        # 规则2：文本位置居中检测
        bbox = block['bbox']
        page_width = block['page_width']
        is_centered = abs((bbox[2] + bbox[0])/2 - page_width/2) < 50  # 左右50px容差
        
        # 规则3：全大写字母检测
        is_uppercase = block['text'].strip().isupper()
        
        # 综合判断：满足任意两个条件即为标题
        conditions = [
            level > 0,
            is_centered,
            is_uppercase
        ]
        return sum(conditions) >= 2, level

    def _convert_excel(self, excel_path, output_dir):
        """处理 Excel 文件"""
        md_content = ["## Excel 表格转换\n"]
        
        xl = pd.ExcelFile(excel_path)
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            md_content.append(f"### 工作表: {sheet_name}")
            md_content.append(df.to_markdown(index=False) + "\n")
        
        self._save_md(md_content, excel_path, output_dir)

    def _save_md(self, content, input_path, output_dir):
        """保存为 Markdown 文件"""
        base_name = os.path.basename(input_path)
        output_path = os.path.join(output_dir, f"{os.path.splitext(base_name)[0]}.md")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(content))
        
        print(f"转换完成 ➜ {output_path}")

# 使用示例
if __name__ == "__main__":
    converter = MarkdownifyMCP()
    converter.convert("docx", "res")
    converter.convert("pdf", "res")
    converter.convert("xlsx", "res")